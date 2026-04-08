from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


def set_run_font(run, size=12, bold=False):
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)


def add_heading(doc, text, level):
    p = doc.add_heading("", level=level)
    r = p.add_run(text)
    set_run_font(r, size=18 if level == 0 else 14, bold=True)
    return p


def add_para(doc, text, bold=False, center=False):
    p = doc.add_paragraph("")
    r = p.add_run(text)
    set_run_font(r, bold=bold)
    if center:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    if p.runs:
        set_run_font(p.runs[0])
    return p


def generate():
    out_path = Path("D:/OCR/docs/hr_archive_function_brief_final_zh_v2.docx")
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(12)

    add_heading(doc, "档案管理服务智能化功能梳理（软件部分）", 0)
    add_para(doc, "汇报定位：第三部分（功能梳理），可直接拆分为 4-5 页 PPT", center=True)
    add_para(doc, "适用对象：人社局业务与管理层", center=True)
    add_para(doc, "")

    add_heading(doc, "一、现状痛点（当前必须解决）", 1)
    add_bullet(doc, "材料类型多：扫描件、照片、PDF并存，人工逐份录入耗时。")
    add_bullet(doc, "版式复杂：表格、印章、手写批注混合，传统识别稳定性不足。")
    add_bullet(doc, "校对成本高：识别后还要二次核对，容易出现漏改、错改。")
    add_bullet(doc, "检索效率低：历史材料查找依赖人工经验，跨批次定位慢。")
    add_bullet(doc, "管理风险点：缺少统一进度视图与质量回看机制。")
    add_para(doc, "结论：系统建设目标不是“炫技术”，而是“提效率、保准确、可追溯、可审计”。", bold=True)

    doc.add_page_break()

    add_heading(doc, "二、第一阶段：已落地功能（现有能力）", 1)
    add_para(doc, "以下内容可按“用例图”方式讲解，避免技术术语堆砌。")

    add_heading(doc, "2.1 参与角色（Actor）", 2)
    add_bullet(doc, "档案专员：导入材料、发起识别、人工校对、归档入库。")
    add_bullet(doc, "审核管理员：复核关键字段、处理冲突项、确认归档质量。")
    add_bullet(doc, "系统管理员：维护运行环境、监控服务可用性。")

    add_heading(doc, "2.2 核心用例（Use Case）", 2)
    add_bullet(doc, "UC-01 批量导入材料（文件/目录）")
    add_bullet(doc, "UC-02 选择识别模式（综合识别/版式识别/文本识别）")
    add_bullet(doc, "UC-03 查看批次进度（总量、处理中、已完成、异常）")
    add_bullet(doc, "UC-04 查看识别结果并人工编辑（文本、表格）")
    add_bullet(doc, "UC-05 保存回写并进入归档检索")
    add_bullet(doc, "UC-06 批次质量概览（冲突项、可核对项）")
    add_bullet(doc, "UC-07 批次问答（带证据片段可追溯）")

    add_heading(doc, "2.3 用例关系（文字图）", 2)
    add_para(doc, "档案专员 -> 批量导入 -> 识别处理 -> 人工校对 -> 保存归档 -> 检索复用")
    add_para(doc, "审核管理员 -> 质量复核 -> 异常处理 -> 结果确认")

    add_heading(doc, "2.4 阶段成效（可现场验证）", 2)
    add_bullet(doc, "批量处理替代逐份手工流程，减少重复性操作。")
    add_bullet(doc, "识别结果支持人工编辑回写，形成“机器识别 + 人工确认”的闭环。")
    add_bullet(doc, "从“找文件”转为“按内容查”，定位速度明显提升。")

    doc.add_page_break()

    add_heading(doc, "三、第一阶段待完善事项（马上要解决）", 1)
    add_para(doc, "这一页建议直接回答“现在还差什么”。")
    add_bullet(doc, "问题1：部分复杂表格存在重复行/重复列输出，需要继续做去重与规则收敛。")
    add_bullet(doc, "问题2：个别结果页编辑后回显不一致，需确保“保存即所见”。")
    add_bullet(doc, "问题3：批次页面异常场景（历史数据清理后）要做自动刷新与防呆提示。")
    add_bullet(doc, "问题4：主界面业务文案仍有技术词，需要进一步统一为政务表达。")
    add_bullet(doc, "问题5：大批量处理时网络波动提示需更友好，降低误报感知。")

    add_heading(doc, "3.1 近期收口动作", 2)
    add_bullet(doc, "统一进度展示：处理未完成不进入正式预览，先看批次进度。")
    add_bullet(doc, "统一保存口径：结构化数据优先展示，原始HTML不覆盖人工修改。")
    add_bullet(doc, "统一错误提示：区分“环境未就绪”和“业务数据为空”。")
    add_bullet(doc, "统一术语体系：面向业务人员，不暴露模型厂商与API实现细节。")

    doc.add_page_break()

    add_heading(doc, "四、第二阶段：AI赋能重点（未来优化）", 1)
    add_para(doc, "定位：在现有可用系统上做“准确优先”的智能增强，不影响主流程稳定。")

    add_heading(doc, "4.1 智能整合（跨材料合并）", 2)
    add_bullet(doc, "按批次自动识别“同一文档的多份材料”，先合并再抽取字段。")
    add_bullet(doc, "输出分组依据与冲突项，支持人工核对，不做黑盒覆盖。")

    add_heading(doc, "4.2 智能字段抽取（双路对比）", 2)
    add_bullet(doc, "规则抽取与模型抽取并行，返回“规则结果/智能结果/推荐结果”。")
    add_bullet(doc, "冲突字段显式提示，确保可解释、可回查。")

    add_heading(doc, "4.3 智能问答（证据可追溯）", 2)
    add_bullet(doc, "基于批次材料做检索增强问答，答案必须带证据片段。")
    add_bullet(doc, "证据不足时明确“无法确认”，避免误导性结论。")

    add_heading(doc, "4.4 质量闭环（持续优化）", 2)
    add_bullet(doc, "引入“人工核对”与反馈机制，沉淀高质量样本。")
    add_bullet(doc, "形成“识别 -> 核对 -> 反馈 -> 优化”的持续提升闭环。")

    doc.add_page_break()

    add_heading(doc, "五、本地化部署与项目落地效果（对甲方口径）", 1)
    add_heading(doc, "5.1 数据安全口径", 2)
    add_bullet(doc, "主链路本地部署：采集、识别、存储、检索均在本地环境完成。")
    add_bullet(doc, "形成数据本地化闭环：原始材料与处理结果不出本地管理域。")
    add_bullet(doc, "外部能力可选增强：可按策略关闭或隔离，不影响本地主流程运行。")

    add_heading(doc, "5.2 项目落地效果（业务可感知）", 2)
    add_bullet(doc, "效率提升：批量处理替代重复录入，缩短材料整理周期。")
    add_bullet(doc, "质量可控：关键信息“机器初提+人工核对”，减少归档差错。")
    add_bullet(doc, "管理可视：批次进度、异常情况、质量状态一屏可见。")
    add_bullet(doc, "追溯可审：每次处理、每次修改、每条证据都可回看。")

    add_heading(doc, "5.3 汇报收尾建议", 2)
    add_para(doc, "建议以“先稳态上线、再智能扩展”为节奏推进：第一阶段保证可用与可控，第二阶段持续做AI赋能，逐步扩大覆盖范围与服务价值。", bold=True)

    doc.save(out_path)
    print(str(out_path))


if __name__ == "__main__":
    generate()
