from datetime import datetime
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_DIR = Path(r"D:\OCR\docs\report_visual_pack")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

SLIDE_W = 1920
SLIDE_H = 1080


def load_font(size: int, bold: bool = False):
    if bold:
        candidates = [
            r"C:\Windows\Fonts\msyhbd.ttc",
            r"C:\Windows\Fonts\simhei.ttf",
            r"C:\Windows\Fonts\msyh.ttc",
        ]
    else:
        candidates = [
            r"C:\Windows\Fonts\msyh.ttc",
            r"C:\Windows\Fonts\simhei.ttf",
            r"C:\Windows\Fonts\msyhbd.ttc",
        ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def split_lines(draw: ImageDraw.ImageDraw, text: str, font, max_width: int):
    lines = []
    for paragraph in text.split("\n"):
        if not paragraph:
            lines.append("")
            continue
        buffer = ""
        for ch in paragraph:
            test = buffer + ch
            bbox = draw.textbbox((0, 0), test, font=font)
            width = bbox[2] - bbox[0]
            if width <= max_width:
                buffer = test
            else:
                if buffer:
                    lines.append(buffer)
                buffer = ch
        if buffer:
            lines.append(buffer)
    return lines


def draw_text_block(
    draw: ImageDraw.ImageDraw,
    text: str,
    box: tuple[int, int, int, int],
    font,
    fill=(35, 49, 70),
    line_gap=8,
):
    x1, y1, x2, y2 = box
    max_width = max(1, x2 - x1)
    lines = split_lines(draw, text, font, max_width)
    y = y1
    for line in lines:
        if y > y2:
            break
        draw.text((x1, y), line, font=font, fill=fill)
        lh = draw.textbbox((0, 0), "行", font=font)[3]
        y += lh + line_gap


def draw_title(draw, title: str, subtitle: str):
    title_font = load_font(56, bold=True)
    sub_font = load_font(28)
    draw.text((80, 38), title, font=title_font, fill=(24, 52, 96))
    draw.text((82, 120), subtitle, font=sub_font, fill=(87, 107, 137))


def rounded_box(draw, box, radius=20, fill=(255, 255, 255), outline=(190, 205, 230), width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def slide1():
    img = Image.new("RGB", (SLIDE_W, SLIDE_H), (244, 247, 252))
    draw = ImageDraw.Draw(img)
    draw_title(
        draw,
        "人社档案系统：现状痛点与建设目标",
        "第三部分（功能梳理）建议讲法：先讲问题，再讲已落地能力",
    )

    left_box = (80, 200, 920, 860)
    right_box = (1000, 200, 1840, 860)
    rounded_box(draw, left_box, fill=(255, 255, 255), outline=(223, 125, 125))
    rounded_box(draw, right_box, fill=(255, 255, 255), outline=(107, 153, 230))

    h_font = load_font(34, bold=True)
    draw.text((120, 238), "当前痛点", font=h_font, fill=(159, 52, 52))
    draw.text((1040, 238), "已落地能力（对应痛点）", font=h_font, fill=(44, 87, 160))

    b_font = load_font(27)
    left_bullets = [
        "• 材料类型杂：纸质扫描、拍照件、PDF并存，人工录入重复。",
        "• 版式复杂：表格、印章、手写批注混合，传统识别稳定性不足。",
        "• 校对压力大：识别后仍要逐条核对，容易漏改、错改。",
        "• 检索效率低：历史档案跨目录、跨批次定位慢。",
        "• 过程不可视：缺少统一进度、异常与质量视图。",
    ]
    right_bullets = [
        "• 支持批量导入（文件/目录）和分阶段进度展示。",
        "• 支持综合识别、版式识别、文本识别三种处理模式。",
        "• 结果页支持文本/表格人工编辑并保存回写。",
        "• 支持历史目录、全文检索、字段提取与归档导出。",
        "• 支持批次质量概览、冲突提示、证据问答与追溯。",
    ]

    draw_text_block(draw, "\n".join(left_bullets), (120, 300, 870, 760), b_font, fill=(85, 41, 41))
    draw_text_block(draw, "\n".join(right_bullets), (1040, 300, 1790, 760), b_font, fill=(37, 65, 111))

    footer_box = (80, 900, 1840, 1010)
    rounded_box(draw, footer_box, radius=16, fill=(235, 243, 255), outline=(137, 170, 221))
    footer_font = load_font(30, bold=True)
    draw.text(
        (110, 935),
        "核心结论：系统价值在于“提效率 + 保质量 + 可追溯”，而不是展示技术名词。",
        font=footer_font,
        fill=(25, 59, 116),
    )
    path = OUTPUT_DIR / "01_现状痛点与建设目标.png"
    img.save(path)
    return path


def draw_actor(draw, x, y, label):
    color = (46, 63, 89)
    draw.ellipse((x - 22, y, x + 22, y + 44), outline=color, width=3)
    draw.line((x, y + 44, x, y + 112), fill=color, width=3)
    draw.line((x - 42, y + 70, x + 42, y + 70), fill=color, width=3)
    draw.line((x, y + 112, x - 32, y + 160), fill=color, width=3)
    draw.line((x, y + 112, x + 32, y + 160), fill=color, width=3)
    font = load_font(30, bold=True)
    tw = draw.textbbox((0, 0), label, font=font)[2]
    draw.text((x - tw / 2, y + 176), label, font=font, fill=color)


def draw_usecase(draw, cx, cy, text, w=290, h=76):
    box = (cx - w // 2, cy - h // 2, cx + w // 2, cy + h // 2)
    rounded_box(draw, box, radius=40, fill=(243, 248, 255), outline=(85, 130, 212), width=3)
    font = load_font(28, bold=True)
    tb = draw.textbbox((0, 0), text, font=font)
    tw = tb[2] - tb[0]
    th = tb[3] - tb[1]
    draw.text((cx - tw / 2, cy - th / 2 - 2), text, font=font, fill=(30, 58, 104))
    return box


def connect_actor_to_usecase(draw, actor_anchor, box, to_left=True):
    x1, y1, x2, y2 = box
    tx = x1 if to_left else x2
    ty = (y1 + y2) // 2
    draw.line((actor_anchor[0], actor_anchor[1], tx, ty), fill=(128, 148, 181), width=2)


def slide2():
    img = Image.new("RGB", (SLIDE_W, SLIDE_H), (244, 247, 252))
    draw = ImageDraw.Draw(img)
    draw_title(
        draw,
        "现有功能用例图（第一阶段已落地）",
        "面向汇报：直接解释“谁在用、怎么用、带来什么效果”",
    )

    boundary = (320, 190, 1600, 940)
    rounded_box(draw, boundary, radius=18, fill=(255, 255, 255), outline=(110, 138, 190), width=3)
    draw.text((350, 218), "系统边界：档案智能整理平台", font=load_font(26, bold=True), fill=(59, 84, 128))

    draw_actor(draw, 120, 330, "档案专员")
    draw_actor(draw, 1800, 330, "审核管理员")

    left_cases = [
        ("批量导入材料", 650, 320),
        ("选择识别模式", 650, 420),
        ("查看批次进度", 650, 520),
        ("人工编辑结果", 650, 620),
        ("保存归档检索", 650, 720),
    ]
    right_cases = [
        ("核对关键信息", 1260, 390),
        ("处理冲突异常", 1260, 510),
        ("确认质量结果", 1260, 630),
    ]
    center_case = ("证据问答追溯", 955, 825, 430)

    left_boxes = [draw_usecase(draw, x, y, t) for t, x, y in left_cases]
    right_boxes = [draw_usecase(draw, x, y, t) for t, x, y in right_cases]
    center_box = draw_usecase(draw, center_case[1], center_case[2], center_case[0], w=center_case[3], h=78)

    for box in left_boxes:
        connect_actor_to_usecase(draw, (175, 400), box, to_left=True)
    connect_actor_to_usecase(draw, (175, 400), center_box, to_left=True)
    for box in right_boxes:
        connect_actor_to_usecase(draw, (1745, 400), box, to_left=False)
    connect_actor_to_usecase(draw, (1745, 400), center_box, to_left=False)

    small = load_font(24)
    draw.text(
        (340, 965),
        "讲解建议：专员侧强调“提效”，管理员侧强调“核准与可追溯”。",
        font=small,
        fill=(96, 113, 137),
    )
    path = OUTPUT_DIR / "02_现有功能用例图.png"
    img.save(path)
    return path


def flow_box(draw, box, title, body, color):
    rounded_box(draw, box, radius=18, fill=(255, 255, 255), outline=color, width=3)
    draw.text((box[0] + 20, box[1] + 18), title, font=load_font(30, bold=True), fill=color)
    draw_text_block(
        draw,
        body,
        (box[0] + 20, box[1] + 66, box[2] - 20, box[3] - 20),
        load_font(24),
        fill=(57, 73, 99),
        line_gap=6,
    )


def slide3():
    img = Image.new("RGB", (SLIDE_W, SLIDE_H), (244, 247, 252))
    draw = ImageDraw.Draw(img)
    draw_title(
        draw,
        "现有系统主流程（从导入到归档）",
        "甲方可看懂版本：流程清晰、职责清晰、数据闭环清晰",
    )

    boxes = [
        (80, 240, 560, 450),
        (640, 240, 1120, 450),
        (1200, 240, 1680, 450),
        (360, 560, 840, 770),
        (960, 560, 1440, 770),
    ]
    flow_box(draw, boxes[0], "1. 材料导入", "文件/目录批量导入\n按批次进入处理队列", (60, 115, 195))
    flow_box(draw, boxes[1], "2. 识别处理", "综合/版式/文本识别\n输出结构化结果", (70, 135, 206))
    flow_box(draw, boxes[2], "3. 人工校对", "文本与表格可编辑\n保存后回写结果", (80, 145, 170))
    flow_box(draw, boxes[3], "4. 归档与检索", "字段提取、Excel导出\n历史目录与全文检索", (81, 141, 98))
    flow_box(draw, boxes[4], "5. 质量与问答", "批次质量概览\n证据问答与反馈闭环", (126, 116, 199))

    # arrows
    arrow = (112, 130, 160)
    draw.line((560, 345, 640, 345), fill=arrow, width=4)
    draw.polygon([(640, 345), (620, 333), (620, 357)], fill=arrow)
    draw.line((1120, 345, 1200, 345), fill=arrow, width=4)
    draw.polygon([(1200, 345), (1180, 333), (1180, 357)], fill=arrow)
    draw.line((1440, 450, 1440, 560), fill=arrow, width=4)
    draw.polygon([(1440, 560), (1428, 540), (1452, 540)], fill=arrow)
    draw.line((960, 665, 840, 665), fill=arrow, width=4)
    draw.polygon([(840, 665), (860, 653), (860, 677)], fill=arrow)

    footer = (80, 860, 1840, 1010)
    rounded_box(draw, footer, radius=16, fill=(235, 243, 255), outline=(140, 171, 221))
    draw_text_block(
        draw,
        "数据安全口径：主链路本地部署，采集-识别-存储-检索均在本地闭环完成；外部增强能力可选且可关闭。",
        (110, 902, 1810, 995),
        load_font(30, bold=True),
        fill=(31, 63, 118),
        line_gap=10,
    )

    path = OUTPUT_DIR / "03_现有系统主流程图.png"
    img.save(path)
    return path


def slide4():
    img = Image.new("RGB", (SLIDE_W, SLIDE_H), (244, 247, 252))
    draw = ImageDraw.Draw(img)
    draw_title(
        draw,
        "第二阶段：AI赋能重点（可落地版本）",
        "讲法：先解决准确率与可解释，再扩展智能服务范围",
    )

    # timeline base
    draw.line((180, 520, 1740, 520), fill=(126, 144, 173), width=5)
    milestones = [
        (320, "阶段A\n同文档整合", "批次内自动识别同一文档\n合并后再抽取字段"),
        (700, "阶段B\n双路抽取比对", "规则+智能并行\n冲突项显式提示"),
        (1080, "阶段C\n质量闭环", "人工核对+反馈\n沉淀高质量样本"),
        (1460, "阶段D\n证据问答", "问答必须带证据\n证据不足则拒答"),
    ]

    for x, title, body in milestones:
        draw.ellipse((x - 16, 504, x + 16, 536), fill=(67, 116, 198), outline=(255, 255, 255), width=2)
        box = (x - 170, 250, x + 170, 470)
        rounded_box(draw, box, radius=18, fill=(255, 255, 255), outline=(114, 152, 213), width=3)
        draw_text_block(draw, title, (box[0] + 18, box[1] + 20, box[2] - 18, box[1] + 92), load_font(30, bold=True), fill=(29, 64, 117), line_gap=4)
        draw_text_block(draw, body, (box[0] + 18, box[1] + 108, box[2] - 18, box[3] - 18), load_font(24), fill=(62, 81, 111), line_gap=6)

    value_box = (180, 620, 1740, 950)
    rounded_box(draw, value_box, radius=18, fill=(255, 255, 255), outline=(164, 180, 209), width=2)
    draw.text((220, 655), "二期落地价值（对甲方可感知）", font=load_font(34, bold=True), fill=(41, 66, 111))
    values = [
        "1) 识别更准：复杂版面与字段冲突可解释、可复核。",
        "2) 管理更稳：批次质量可量化，问题可定位、可追责。",
        "3) 服务更广：从“识别工具”升级为“档案智能服务平台”。",
    ]
    draw_text_block(draw, "\n".join(values), (230, 720, 1680, 920), load_font(30), fill=(53, 74, 107), line_gap=12)

    path = OUTPUT_DIR / "04_二期AI赋能路线图.png"
    img.save(path)
    return path


def slide5():
    img = Image.new("RGB", (SLIDE_W, SLIDE_H), (244, 247, 252))
    draw = ImageDraw.Draw(img)
    draw_title(
        draw,
        "项目落地效果与实施建议",
        "8分钟汇报建议收束页：讲效果、讲计划、讲需要甲方确认什么",
    )

    cards = [
        ((80, 230, 900, 430), "效率效果", ["批量替代逐份处理", "进度可视化，减少等待盲区", "从“找文件”转向“按内容查找”"], (67, 127, 198)),
        ((1020, 230, 1840, 430), "质量效果", ["识别+人工核对形成闭环", "冲突字段可追溯可解释", "批次质量可持续改进"], (74, 151, 126)),
        ((80, 470, 900, 670), "安全效果", ["主链路本地部署、本地存储", "原始材料不出本地管理域", "外部增强能力可选可关闭"], (122, 113, 186)),
        ((1020, 470, 1840, 670), "管理效果", ["历史记录可回看可追责", "异常项有统一处理入口", "支持扩面复制到更多业务线"], (188, 124, 71)),
    ]

    for box, title, points, color in cards:
        rounded_box(draw, box, radius=18, fill=(255, 255, 255), outline=color, width=3)
        draw.text((box[0] + 22, box[1] + 18), title, font=load_font(34, bold=True), fill=color)
        draw_text_block(
            draw,
            "\n".join([f"• {item}" for item in points]),
            (box[0] + 24, box[1] + 75, box[2] - 24, box[3] - 20),
            load_font(26),
            fill=(58, 76, 106),
            line_gap=8,
        )

    step_box = (80, 720, 1840, 970)
    rounded_box(draw, step_box, radius=18, fill=(233, 242, 255), outline=(130, 166, 223), width=2)
    draw.text((110, 752), "建议实施节奏（稳妥上线）", font=load_font(34, bold=True), fill=(26, 58, 110))
    steps = "试点验证（1个部门）  ->  扩面上线（多部门）  ->  稳态运营（持续AI优化）"
    draw.text((110, 820), steps, font=load_font(32, bold=True), fill=(34, 71, 130))
    draw.text(
        (110, 882),
        "需甲方确认：试点范围、验收口径、推广节奏、运维责任边界。",
        font=load_font(28),
        fill=(64, 84, 118),
    )

    path = OUTPUT_DIR / "05_落地效果与实施建议.png"
    img.save(path)
    return path


def build_doc(image_paths: list[Path]):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(12)

    title = doc.add_paragraph("人社档案系统功能梳理与AI赋能建议（图文版）")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(20)
    title.runs[0].font.name = "Microsoft YaHei"
    title.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    info = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sections = [
        ("第1页：现状痛点与建设目标", "建议讲45秒：先讲痛点，再讲已落地能力，形成“问题-能力”对应。"),
        ("第2页：现有功能用例图", "建议讲60秒：按角色讲，不讲底层技术。专员侧讲效率，管理员侧讲质量。"),
        ("第3页：现有系统主流程", "建议讲60秒：导入-识别-校对-归档-质量，强调本地数据闭环。"),
        ("第4页：二期AI赋能路线", "建议讲90秒：明确“先准确、再扩展”的节奏，避免空泛。"),
        ("第5页：落地效果与实施建议", "建议讲60秒：讲已见效果、可见收益、甲方需确认项。"),
    ]

    for idx, (img_path, sec) in enumerate(zip(image_paths, sections), start=1):
        doc.add_page_break()
        h = doc.add_paragraph(sec[0])
        h.runs[0].bold = True
        h.runs[0].font.size = Pt(14)
        h.runs[0].font.name = "Microsoft YaHei"
        h.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        doc.add_paragraph(sec[1])
        doc.add_picture(str(img_path), width=Inches(9.5))

    out_doc = OUTPUT_DIR / "人社档案系统_功能梳理与AI规划_图文版.docx"
    doc.save(out_doc)
    return out_doc


def main():
    images = [slide1(), slide2(), slide3(), slide4(), slide5()]
    doc = build_doc(images)
    print("IMAGES:")
    for image in images:
        print(image)
    print("DOC:")
    print(doc)


if __name__ == "__main__":
    main()
