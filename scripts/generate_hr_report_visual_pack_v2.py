from datetime import datetime
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SLIDE_W = 1920
SLIDE_H = 1080
OUTPUT_DIR = Path(r"D:\OCR\docs\report_visual_pack_v2")


def ensure_dir():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def load_font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\msyh.ttc",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def rounded(draw: ImageDraw.ImageDraw, box, radius=20, fill=(255, 255, 255), outline=(180, 200, 230), width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def split_lines(draw: ImageDraw.ImageDraw, text: str, font, max_width: int):
    lines: list[str] = []
    for paragraph in text.split("\n"):
        if not paragraph:
            lines.append("")
            continue
        current = ""
        for ch in paragraph:
            test = current + ch
            width = draw.textbbox((0, 0), test, font=font)[2]
            if width <= max_width:
                current = test
            else:
                if current:
                    lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines


def draw_block(
    draw: ImageDraw.ImageDraw,
    text: str,
    box: tuple[int, int, int, int],
    font,
    fill=(40, 58, 83),
    line_gap=8,
):
    x1, y1, x2, y2 = box
    lines = split_lines(draw, text, font, max(1, x2 - x1))
    y = y1
    for line in lines:
        if y > y2:
            break
        draw.text((x1, y), line, font=font, fill=fill)
        h = draw.textbbox((0, 0), "字", font=font)[3]
        y += h + line_gap


def draw_title(draw: ImageDraw.ImageDraw, title: str, subtitle: str):
    draw.text((70, 32), title, font=load_font(56, bold=True), fill=(23, 51, 96))
    draw.text((72, 116), subtitle, font=load_font(28), fill=(86, 106, 135))


def slide1_pain_vs_goal():
    img = Image.new("RGB", (SLIDE_W, SLIDE_H), (244, 247, 252))
    draw = ImageDraw.Draw(img)
    draw_title(draw, "人社档案系统：现状痛点与建设目标", "第三部分建议讲法：痛点 -> 对应能力 -> 结论")

    left = (70, 190, 930, 870)
    right = (990, 190, 1850, 870)
    rounded(draw, left, fill=(255, 255, 255), outline=(223, 132, 132))
    rounded(draw, right, fill=(255, 255, 255), outline=(110, 154, 228))

    draw.text((110, 230), "当前痛点", font=load_font(36, bold=True), fill=(150, 53, 53))
    draw.text((1030, 230), "对应建设目标", font=load_font(36, bold=True), fill=(44, 87, 160))

    pain_text = "\n".join(
        [
            "• 材料类型杂：扫描件、照片、PDF并存，人工录入重复。",
            "• 版式复杂：表格、印章、手写批注混合，稳定性不足。",
            "• 校对压力大：识别后仍要逐条核对，容易漏改错改。",
            "• 检索效率低：历史档案跨目录跨批次定位慢。",
            "• 过程不可视：缺少统一进度、异常、质量视图。",
        ]
    )
    goal_text = "\n".join(
        [
            "• 建立批量导入和分阶段处理机制，减少重复劳动。",
            "• 构建多模式识别能力，覆盖复杂版式与常规文本。",
            "• 建立“机器识别 + 人工复核”闭环，保障结果可控。",
            "• 实现按内容检索与按批次回看，提升调用效率。",
            "• 建立质量概览与证据问答，支持可追溯管理。",
        ]
    )
    draw_block(draw, pain_text, (110, 300, 880, 780), load_font(32), fill=(87, 46, 46), line_gap=10)
    draw_block(draw, goal_text, (1030, 300, 1800, 780), load_font(32), fill=(36, 64, 112), line_gap=10)

    footer = (70, 900, 1850, 1025)
    rounded(draw, footer, radius=16, fill=(234, 242, 255), outline=(136, 170, 221))
    draw_block(
        draw,
        "结论：本项目核心价值是“提效率、保质量、可追溯”，并形成可持续优化基础。",
        (100, 938, 1810, 1000),
        load_font(34, bold=True),
        fill=(26, 58, 110),
        line_gap=8,
    )

    out = OUTPUT_DIR / "01_现状痛点与建设目标.png"
    img.save(out)
    return out


def draw_actor(draw: ImageDraw.ImageDraw, x: int, y: int, label: str):
    color = (42, 61, 87)
    draw.ellipse((x - 20, y, x + 20, y + 40), outline=color, width=3)
    draw.line((x, y + 40, x, y + 105), fill=color, width=3)
    draw.line((x - 40, y + 65, x + 40, y + 65), fill=color, width=3)
    draw.line((x, y + 105, x - 30, y + 150), fill=color, width=3)
    draw.line((x, y + 105, x + 30, y + 150), fill=color, width=3)
    tw = draw.textbbox((0, 0), label, font=load_font(28, bold=True))[2]
    draw.text((x - tw / 2, y + 165), label, font=load_font(28, bold=True), fill=color)


def draw_usecase(draw: ImageDraw.ImageDraw, cx: int, cy: int, text: str, w=300, h=76):
    box = (cx - w // 2, cy - h // 2, cx + w // 2, cy + h // 2)
    rounded(draw, box, radius=38, fill=(245, 249, 255), outline=(86, 131, 214), width=3)
    tb = draw.textbbox((0, 0), text, font=load_font(29, bold=True))
    tw = tb[2] - tb[0]
    th = tb[3] - tb[1]
    draw.text((cx - tw / 2, cy - th / 2 - 2), text, font=load_font(29, bold=True), fill=(30, 58, 104))
    return box


def connect(draw: ImageDraw.ImageDraw, start: tuple[int, int], box, to_left: bool):
    x1, y1, x2, y2 = box
    tx = x1 if to_left else x2
    ty = (y1 + y2) // 2
    draw.line((start[0], start[1], tx, ty), fill=(130, 150, 185), width=2)


def slide2_usecase():
    img = Image.new("RGB", (SLIDE_W, SLIDE_H), (244, 247, 252))
    draw = ImageDraw.Draw(img)
    draw_title(draw, "第一阶段功能用例图（已落地能力）", "按角色讲流程：档案专员 / 审核管理员 / 系统管理员")

    boundary = (280, 180, 1640, 930)
    rounded(draw, boundary, radius=20, fill=(255, 255, 255), outline=(109, 139, 193), width=3)
    draw.text((320, 212), "系统边界：档案智能整理平台", font=load_font(28, bold=True), fill=(57, 84, 129))

    draw_actor(draw, 110, 310, "档案专员")
    draw_actor(draw, 1810, 310, "审核管理员")
    draw_actor(draw, 960, 930, "系统管理员")

    left = [("批量导入材料", 660, 310), ("选择识别模式", 660, 410), ("查看批次进度", 660, 510), ("编辑与保存结果", 660, 610)]
    right = [("核对关键信息", 1270, 360), ("处理冲突异常", 1270, 500), ("确认质量结论", 1270, 640)]
    center = ("证据问答追溯", 960, 780, 430)

    left_boxes = [draw_usecase(draw, x, y, t) for t, x, y in left]
    right_boxes = [draw_usecase(draw, x, y, t) for t, x, y in right]
    center_box = draw_usecase(draw, center[1], center[2], center[0], w=center[3], h=78)

    for b in left_boxes:
        connect(draw, (165, 375), b, to_left=True)
    connect(draw, (165, 375), center_box, to_left=True)
    for b in right_boxes:
        connect(draw, (1755, 375), b, to_left=False)
    connect(draw, (1755, 375), center_box, to_left=False)
    connect(draw, (960, 930), center_box, to_left=False)

    draw.text(
        (300, 965),
        "讲解提示：专员侧强调“提效”，管理员侧强调“核准与可追溯”，系统管理员侧强调“稳定与审计”。",
        font=load_font(24),
        fill=(96, 113, 137),
    )

    out = OUTPUT_DIR / "02_第一阶段功能用例图.png"
    img.save(out)
    return out


def matrix_cell(draw: ImageDraw.ImageDraw, box, title: str, current: str, near: str, next_stage: str):
    rounded(draw, box, radius=14, fill=(255, 255, 255), outline=(138, 165, 209), width=2)
    draw.text((box[0] + 14, box[1] + 10), title, font=load_font(24, bold=True), fill=(36, 68, 122))
    draw_block(draw, f"当前：{current}", (box[0] + 14, box[1] + 48, box[2] - 14, box[1] + 112), load_font(20), fill=(67, 83, 110), line_gap=4)
    draw_block(draw, f"近期：{near}", (box[0] + 14, box[1] + 112, box[2] - 14, box[1] + 176), load_font(20), fill=(40, 111, 86), line_gap=4)
    draw_block(draw, f"二期：{next_stage}", (box[0] + 14, box[1] + 176, box[2] - 14, box[3] - 12), load_font(20), fill=(123, 95, 49), line_gap=4)


def slide3_goal_matrix():
    img = Image.new("RGB", (SLIDE_W, SLIDE_H), (244, 247, 252))
    draw = ImageDraw.Draw(img)
    draw_title(draw, "六项目标落地矩阵（现状 / 近期上线 / 二期优化）", "把你的6个目标逐条落到执行路径，避免空话")

    cells = [
        ("目标1 双层PDF导出", "支持结果导出与文件预览", "新增双层PDF导出（影像层+文本层）", "支持批量合并导出与目录化下载"),
        ("目标2 字段自动归档", "已支持规则字段提取与编辑", "支持手动设字段 + 智能抽取", "输出置信度与人工复审建议"),
        ("目标3 分类保存", "支持按批次管理与历史目录", "支持按批次/目录归档保存", "支持按内容主题自动分类归档"),
        ("目标4 自动整合", "已支持批次处理与结果聚合", "上线同文档自动归并", "给出可追溯分组依据与冲突说明"),
        ("目标5 缺页提醒", "可查看页数与处理状态", "增加页码连续性检测", "增加内容完整性评分与复查提醒"),
        ("目标6 对话整理", "已支持批次问答入口", "支持选定文档问答整理", "支持一批文档自动汇总与证据链输出"),
    ]

    start_x = 70
    start_y = 190
    w = 580
    h = 250
    gap_x = 45
    gap_y = 22
    for i, (title, current, near, nxt) in enumerate(cells):
        row = i // 3
        col = i % 3
        x1 = start_x + col * (w + gap_x)
        y1 = start_y + row * (h + gap_y)
        box = (x1, y1, x1 + w, y1 + h)
        matrix_cell(draw, box, title, current, near, nxt)

    footer = (70, 960, 1850, 1030)
    rounded(draw, footer, radius=14, fill=(235, 243, 255), outline=(138, 170, 220), width=2)
    draw_block(
        draw,
        "结论：六项目标中，前4项可作为第一阶段增强交付，后2项作为二期AI重点推进。",
        (95, 978, 1830, 1018),
        load_font(28, bold=True),
        fill=(25, 60, 116),
        line_gap=4,
    )

    out = OUTPUT_DIR / "03_六项目标落地矩阵.png"
    img.save(out)
    return out


def stage_card(draw: ImageDraw.ImageDraw, box, title: str, body: str):
    rounded(draw, box, radius=18, fill=(255, 255, 255), outline=(111, 151, 214), width=3)
    draw_block(draw, title, (box[0] + 16, box[1] + 14, box[2] - 16, box[1] + 80), load_font(42, bold=True), fill=(30, 62, 117), line_gap=4)
    draw_block(draw, body, (box[0] + 16, box[1] + 86, box[2] - 16, box[3] - 16), load_font(31), fill=(60, 81, 112), line_gap=8)


def slide4_ai_roadmap():
    img = Image.new("RGB", (SLIDE_W, SLIDE_H), (244, 247, 252))
    draw = ImageDraw.Draw(img)
    draw_title(draw, "第二阶段AI赋能路线图（A/B/C/D）", "推进原则：准确优先、可解释优先、可追溯优先")

    cards = [
        ((110, 200, 500, 470), "阶段A", "同文档整合\n自动识别同一文档\n先合并再抽取"),
        ((560, 200, 950, 470), "阶段B", "双路抽取比对\n规则+智能并行\n冲突显式提示"),
        ((1010, 200, 1400, 470), "阶段C", "质量闭环\n人工核对+反馈\n沉淀高质量样本"),
        ((1460, 200, 1850, 470), "阶段D", "证据问答\n答案必须带证据\n证据不足拒答"),
    ]
    for box, t, body in cards:
        stage_card(draw, box, t, body)

    draw.line((180, 560, 1740, 560), fill=(119, 140, 173), width=5)
    for x in (305, 755, 1205, 1655):
        draw.ellipse((x - 14, 546, x + 14, 574), fill=(65, 115, 198), outline=(255, 255, 255), width=2)

    panel = (110, 650, 1850, 970)
    rounded(draw, panel, radius=16, fill=(255, 255, 255), outline=(166, 181, 208), width=2)
    draw.text((150, 688), "二期可感知效果", font=load_font(44, bold=True), fill=(43, 70, 116))
    draw_block(
        draw,
        "1) 识别更准：复杂版面、缺页和冲突项可解释、可复核。\n"
        "2) 管理更稳：质量结果可量化，问题可定位、可追责。\n"
        "3) 服务更广：从“识别工具”升级到“档案智能服务平台”。",
        (150, 760, 1780, 935),
        load_font(34),
        fill=(53, 74, 106),
        line_gap=12,
    )

    out = OUTPUT_DIR / "04_第二阶段AI赋能路线图.png"
    img.save(out)
    return out


def slide5_effect_and_plan():
    img = Image.new("RGB", (SLIDE_W, SLIDE_H), (244, 247, 252))
    draw = ImageDraw.Draw(img)
    draw_title(draw, "落地效果与实施建议（决策页）", "讲法：先说效果，再说实施节奏，再明确甲方确认项")

    cards = [
        ((70, 210, 930, 430), "效率效果", ["批量替代逐份处理", "进度可视化减少等待盲区", "从“找文件”转为“按内容查找”"], (62, 121, 197)),
        ((990, 210, 1850, 430), "质量效果", ["识别+人工复核形成闭环", "字段冲突可解释可追溯", "批次质量可持续优化"], (69, 149, 121)),
        ((70, 470, 930, 690), "安全效果", ["主链路本地部署本地存储", "原始材料不出本地管理域", "外部增强可选可关闭"], (118, 108, 185)),
        ((990, 470, 1850, 690), "管理效果", ["历史记录可回看可追责", "异常项有统一处理入口", "支持扩面复制到更多业务线"], (184, 121, 67)),
    ]

    for box, title, points, color in cards:
        rounded(draw, box, radius=18, fill=(255, 255, 255), outline=color, width=3)
        draw.text((box[0] + 22, box[1] + 16), title, font=load_font(45, bold=True), fill=color)
        draw_block(draw, "\n".join([f"• {p}" for p in points]), (box[0] + 22, box[1] + 78, box[2] - 22, box[3] - 18), load_font(34), fill=(58, 76, 106), line_gap=12)

    step = (70, 760, 1850, 1015)
    rounded(draw, step, radius=16, fill=(234, 242, 255), outline=(132, 168, 223), width=2)
    draw.text((100, 802), "建议实施节奏（稳妥上线）", font=load_font(46, bold=True), fill=(27, 58, 110))
    draw.text(
        (100, 874),
        "试点验证（1个部门）  ->  扩面上线（多部门）  ->  稳态运营（持续AI优化）",
        font=load_font(42, bold=True),
        fill=(34, 71, 130),
    )
    draw.text((100, 947), "需甲方确认：试点范围、验收口径、推广节奏、运维责任边界。", font=load_font(33), fill=(66, 85, 118))

    out = OUTPUT_DIR / "05_落地效果与实施建议.png"
    img.save(out)
    return out


def set_run_font(run, size=12, bold=False):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_title_para(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, size=16, bold=True)
    return p


def add_note(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=11, bold=False)
    return p


def add_bullets(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(line)
        set_run_font(r, size=11, bold=False)


def build_word(image_paths: list[Path]):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(11)

    page_data = [
        (
            "第1页：现状痛点与建设目标",
            [
                "现状问题聚焦在“效率、准确、可视化”三方面。",
                "已落地能力与痛点一一对应，可现场验证。",
                "结论是先稳定交付，再持续优化。",
            ],
            "一句话结论：系统建设目标是提效率、保质量、可追溯。",
        ),
        (
            "第2页：第一阶段功能用例图",
            [
                "档案专员负责导入、识别、编辑、归档。",
                "审核管理员负责核对冲突、确认质量。",
                "系统管理员保障平台稳定与审计可用。",
            ],
            "一句话结论：第一阶段已形成“识别-复核-归档”的业务闭环。",
        ),
        (
            "第3页：六项目标落地矩阵",
            [
                "6项目标逐条拆解为“当前状态 / 近期上线 / 二期优化”。",
                "第一阶段优先交付：双层PDF、字段归档、分类保存、自动整合。",
                "第二阶段重点推进：缺页检测、对话整理与证据问答。",
            ],
            "一句话结论：目标分层清晰，避免一次性摊大饼。",
        ),
        (
            "第4页：第二阶段AI赋能路线图",
            [
                "阶段A/B/C/D分别解决整合、比对、闭环、问答。",
                "坚持“准确优先、可解释优先、可追溯优先”。",
                "每阶段都有可验收输出，不依赖空泛指标。",
            ],
            "一句话结论：二期AI建设有路径、有边界、可落地。",
        ),
        (
            "第5页：落地效果与实施建议",
            [
                "从效率、质量、安全、管理四维展示落地价值。",
                "实施建议采用“试点 -> 扩面 -> 稳态运营”节奏。",
                "明确甲方需确认的范围、口径、节奏和责任边界。",
            ],
            "一句话结论：该方案可稳妥上线并支持后续扩展。",
        ),
    ]

    for idx, (title, points, conclusion) in enumerate(page_data):
        if idx > 0:
            doc.add_page_break()
        add_title_para(doc, title)
        add_note(doc, f"讲解提示：{datetime.now().strftime('%Y-%m-%d')} 版本，按45-90秒/页讲解。")
        add_bullets(doc, points)
        doc.add_picture(str(image_paths[idx]), width=Inches(8.6))
        add_note(doc, conclusion)

    doc.add_page_break()
    add_title_para(doc, "附录：能力口径定义表（6项目标）")
    add_note(doc, "说明：用于统一汇报口径，不是接口文档。")
    table = doc.add_table(rows=1, cols=5)
    headers = ["能力项", "输入", "输出", "复核方式", "落地状态"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
    rows = [
        ("双层PDF导出", "识别结果+原始页", "可下载双层PDF", "抽检可搜索文本", "一期近期上线"),
        ("字段自动归档", "材料全文+字段配置", "字段值+置信度", "建议人工复审", "一期增强"),
        ("分类保存", "批次信息+内容特征", "批次/主题目录", "人工抽样复核", "一期增强"),
        ("自动整合", "分散文件集合", "同文档合并结果", "分组依据可追溯", "一期增强"),
        ("缺页提醒", "页序+内容覆盖", "缺页/不全提示", "人工复查确认", "二期重点"),
        ("对话整理", "批次或指定文档", "答案+证据片段", "证据不足拒答", "二期重点"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    out_cn = OUTPUT_DIR / "人社档案系统_功能梳理_5页图文版_V2.docx"
    out_en = OUTPUT_DIR / "hr_archive_report_5page_visual_v2.docx"
    doc.save(out_cn)
    doc.save(out_en)
    return out_cn, out_en


def main():
    ensure_dir()
    images = [
        slide1_pain_vs_goal(),
        slide2_usecase(),
        slide3_goal_matrix(),
        slide4_ai_roadmap(),
        slide5_effect_and_plan(),
    ]
    out_cn, out_en = build_word(images)

    print("IMAGES:")
    for path in images:
        print(path)
    print("WORD_CN:")
    print(out_cn)
    print("WORD_EN:")
    print(out_en)


if __name__ == "__main__":
    main()
